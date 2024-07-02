from docx import Document as DocxDocument
from typing import Iterator
from langchain_core.documents import Document
def lazy_load(file_path) -> Iterator[Document]:  # <-- Does not take any arguments
        """A lazy loader that reads a file line by line.

        When you're implementing lazy load methods, you should use a generator
        to yield documents one by one.
        """
        try:
            doc = DocxDocument(file_path)
            line_number = 0
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    yield Document(
                        page_content=para.text,
                        metadata={"line_number": line_number, "source": file_path},
                    )
                    line_number += 1
        except Exception as e:
            print(f"Error reading {file_path}: {str(e)}")